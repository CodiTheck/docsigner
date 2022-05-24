"""
Doc Signer
==========
Program to sign a document from an image file inserted in it.

@author Dr Mokira
@date   2022-22-05

"""
import os
import enum
from PIL import Image
from pdfrw import PdfReader, PdfWriter, PageMerge
from docx import Document
from docx.shared import Cm

__version__ = '0.0.1';
__author__  = 'Dr Mokira';


class dsg:
    # Python program to print
    # colored text and background
    class color:
        """ Colors class:reset all colors with colors.reset; two
            sub classes fg for foreground
            and bg for background; use as colors.subclass.colorname.
            i.e. colors.fg.red or colors.bg.greenalso, the generic bold, disable,
            underline, reverse, strike through,
            and invisible work with the main class i.e. colors.bold """

        RESET           = '\033[0m';
        BOLD            = '\033[01m';
        DISABLE         = '\033[02m';
        UNDERLINE       = '\033[04m';
        REVERSE         = '\033[07m';
        STRIKETHROUGH   = '\033[09m';
        INVISIBLE       = '\033[08m';

        class FG:
            BLACK       = '\033[30m';
            RED         = '\033[31m';
            GREEN       = '\033[32m';
            ORANGE      = '\033[33m';
            BLUE        = '\033[34m';
            PURPLE      = '\033[35m';
            CYAN        = '\033[36m';
            LIGHTGREY   = '\033[37m';
            DARKGREY    = '\033[90m';
            LIGHTRED    = '\033[91m';
            LIGHTGREEN  = '\033[92m';
            YELLOW      = '\033[93m';
            LIGHTBLUE   = '\033[94m';
            PINK        = '\033[95m';
            LIGHTCYAN   = '\033[96m';

        class BG:
            BLACK       = '\033[40m';
            RED         = '\033[41m';
            GREEN       = '\033[42m';
            ORANGE      = '\033[43m';
            BLUE        = '\033[44m';
            PURPLE      = '\033[45m';
            CYAN        = '\033[46m';
            LIGHTGREY   = '\033[47m';


    @staticmethod
    def log(cl, type, message):
        """ Function to make log in terminal. """
        print("{col}{fgc} {type} {reset} \t{message}".format(col=cl, fgc=dsg.color.FG.BLACK, type=type, reset=dsg.color.RESET, message=message));


    @staticmethod
    def printinfo(message):
        """ Function that is used to print infos in terminal. """
        dsg.log(dsg.color.BG.LIGHTGREY, 'INFO', message);


    @staticmethod
    def printwarn(message):
        """ Function that is used to print warnings in terminal. """
        dsg.log(dsg.color.BG.ORANGE, 'WARN', message);


    @staticmethod
    def printerr(message):
        """ Function that is used to print errors in terminal. """
        dsg.log(dsg.color.BG.RED, 'ERRO', message);

    @staticmethod
    def printsucc(message):
        """ Function that is used to print success message in terminal. """
        dsg.log(dsg.color.BG.GREEN, 'SUCC', message);


    @staticmethod
    def print_err_message(message):
        return "{col}{fgc} {message} {reset} \t".format(col=dsg.color.BG.RED, fgc=dsg.color.FG.BLACK, reset=dsg.color.RESET, message=message);


#######################################################################################################

    class DocType(enum.Enum):
        WORD = 0x001;
        PDF  = 0x002;


    class RelativePos(enum.Enum):
        BOTTOM_CENTER = 0x010;
        BOTTOM_RIGHT  = 0x011;
        BOTTOM_LEFT   = 0x012;
        TOP_CENTER    = 0x020;
        TOP_RIGHT     = 0x021;
        TOP_LEFT      = 0x022;

        CENTER_RIGHT  = 0x001;
        CENTER_LEFT   = 0x002;
        CENTER        = 0x000;

        def abspos(self, recsize, objsize=(0, 0)):
            rec_w = recsize[0];
            rec_h = recsize[1];
            obj_w = objsize[0];
            obj_h = objsize[1];
            x = 0;
            y = 0;

            if self.value == self.BOTTOM_CENTER.value:
                x = int(rec_w / 2) - int(obj_w / 2);
                y = 0;
            elif self.value == self.BOTTOM_RIGHT.value:
                x = rec_w - obj_w;
                y = 0;
            elif self.value == self.BOTTOM_LEFT.value:
                x = 0;
                y = 0;
            elif self.value == self.TOP_CENTER.value:
                x = int(rec_w / 2) - int(obj_w / 2);
                y = rec_h - obj_h;
            elif self.value == self.TOP_RIGHT.value:
                x = rec_w - obj_w;
                y = rec_h - obj_h;
            elif self.value == self.TOP_LEFT.value:
                x = 0;
                y = rec_h - obj_h;
            elif self.value == self.CENTER_RIGHT.value:
                x = rec_w - obj_w;
                y = int(rec_h / 2) - int(obj_h / 2);
            elif self.value == self.CENTER_LEFT.value:
                x = 0;
                y = int(rec_h / 2) - int(obj_h / 2);
            elif self.value == self.CENTER.value:
                x = int(rec_w / 2) - int(obj_w / 2);
                y = int(rec_h / 2) - int(obj_h / 2);

            return x, y;

        def getaxis(self):
            if self.value == self.BOTTOM_CENTER.value:   return (1, 1);
            elif self.value == self.BOTTOM_RIGHT.value:  return (-1, 1);
            elif self.value == self.BOTTOM_LEFT.value:   return (1, 1);
            elif self.value == self.TOP_CENTER.value:    return (1, -1);
            elif self.value == self.TOP_RIGHT.value:     return (-1, -1);
            elif self.value == self.TOP_LEFT.value:      return (1, -1);
            elif self.value == self.CENTER_RIGHT.value:  return (-1, 1);
            elif self.value == self.CENTER_LEFT.value:   return (1, 1);
            elif self.value == self.CENTER.value:        return (1, 1);


    @staticmethod
    def sign(din, dout, simg, doctype=None, pgn=-1, sdim=(128, 128), pos=(0, 0), rlpos=None):
        """ Function that is used to sign a document.
            :args:
                din  [string] represents the location of the document to be signed.
                dout [string] represents the location to the signed document obtained after signing.
                simg [string] represents the image of the signature.
                doctype [dsg.DocType] Represents the type of document you want to sign.
                sdim   [tuple] represents the dimensions of the signature image.
                pos    [tuple] represents the position of signature image on document page
                rlpos  [dsg.RelativePos] represents the relative position of signature image on document page.

            :return:
                True, if the signing operation is successful,
                False, else.
        """
        assert type(din)  is str, dsg.print_err_message("[din]  variable must be a string type.");
        assert type(dout) is str, dsg.print_err_message("[dout] variable must be a string type.");
        assert type(simg) is str, dsg.print_err_message("[simg] variable must be a string type.");
        assert doctype is None or type(doctype) is dsg.DocType, dsg.print_err_message("[doctype] variable must be a dsg.DocType type or None.");
        assert rlpos is None or type(rlpos) is dsg.RelativePos, dsg.print_err_message("[rlpos] variable must be a dsg.RelativePos type or None.");
        assert type(sdim) is tuple and len(sdim) == 2, dsg.print_err_message("[sdim] variable must be a tuple type with two elements (width, height).");
        assert type(pos) is tuple and len(pos) == 2, dsg.print_err_message("[pos] variable must be a tuple type with two elements (x, y).");
        assert type(pgn) is int, dsg.print_err_message("[sdim] variable must be an integer type.");

        resp = None;
        if doctype == dsg.DocType.PDF:
            resp = dsg._sign_pdf(din, dout, simg, pgn, sdim, rlpos, pos);
        elif doctype == dsg.DocType.WORD:
            resp = dsg._sign_docx(din, dout, simg, pgn, sdim);
            # dsg.printwarn(f"The [pos] function is not yet supported for this type of document.");
        else:
            dsg.printinfo("You must specify the type of document to be signed.");
            return False;

        if resp == True:
            dsg.printsucc(f"{din} is signed to {dout} successfully !");
            return True;
        else:
            dsg.printerr(f"{resp}");
            return False;


    @staticmethod
    def __adjust(p1, p2, x=36, y=36):
        info2 = PageMerge().add(p2);
        x2, y2, w2, h2 = info2.xobj_box;
        x += w2;
        y += h2;
        viewrect = (x, y, (w2 - x2 - 2 * x), (h2 - y2 - 2 * y));
        page     = PageMerge(p1).add(p2, viewrect=viewrect);
        return page.render();


    @staticmethod
    def __getrealpos(pos, origin=(0, 0), axis=(1, 1)):
        org_x = origin[0];
        org_y = origin[1];
        axi_x = axis[0];
        axi_y = axis[1];
        return ((org_x + axi_x * pos[0]), (org_y + axi_y * pos[1]));


    @staticmethod
    def _sign_pdf(din, dout, simg, pgn, sdim, rlpos, pos):
        try:
            # We convert the signature to pdf doc, in first.
            img1 = Image.open(simg);
            im1 = img1.resize(sdim);
            im1 = im1.convert("RGB");
            im1.save('.sign.pdf', quality=95);

            watermark_file = '.sign.pdf';
            input_file     = din;
            output_file    = dout;

            # define the reader and writer objects
            reader_input  = PdfReader(input_file)
            writer_output = PdfWriter();

            # we open signature doc.
            watermark_input = PdfReader(watermark_file);
            watermark       = watermark_input.pages[0];
            wat_x, wat_y, wat_w, wat_h = (PageMerge().add(watermark)).xobj_box;

            # insert the signature image on the doc selected page.
            cp = len(reader_input.pages);
            pn = (cp - 1) if (pgn == -1 or pgn < 1 or pgn > cp) else (pgn - 1);

            page = reader_input.pages[pn];
            page_x, page_y, page_w, page_h = (PageMerge().add(page)).xobj_box;

            # building of positon of image
            axis = (1, 1);
            x = 0;
            y = 0;
            print(rlpos);
            if rlpos:
                x, y = rlpos.abspos(recsize=(page_w, page_h), objsize=(wat_w, wat_h));
                axis = rlpos.getaxis();

            x, y = dsg.__getrealpos(pos, (x, y), axis);

            # merger = PageMerge(reader_input.pages[pn]);
            # merger.add(watermark).render();
            # ajust image with buided position
            dsg.__adjust(page, watermark, x, y);

            # write the modified content to disk
            writer_output.write(output_file, reader_input);
            if os.path.exists(watermark_file):
                # if the swap file is exists, delete it
                os.remove(watermark_file);

            return True;
        except Exception as e:
            return e.args[0];


    @staticmethod
    def _sign_docx(din, dout, simg, pgn, sdim):
        try:
            # we open doc
            document = Document(din);

            # we put 3 new paragraphs for ower signature
            p = document.add_paragraph();
            p = document.add_paragraph();
            p = document.add_paragraph();
            r = p.add_run();

            # we put the signature and save the data doc
            px = lambda cm: cm / 37.795275591;
            r.add_picture(simg, width=Cm(px(sdim[0])), height=Cm(px(sdim[1])));

            document.save(dout);
            return True;
        except Exception as e:
            return e.args[0];


if __name__ == '__main__':
    import argparse

    # command line arguments configuration
    parser = argparse.ArgumentParser(description="Program that use to merge an image to document page.");
    parser.add_argument('-t', '--type', default='pdf', dest='doctype', help='Document type.', choices=('pdf', 'word'), required=True);
    parser.add_argument('-i', '--in',   default='./input.pdf', dest='din', help='Input document.', type=str, required=True);
    parser.add_argument('-n', '--page-number', default=-1, dest='page', help='Page number.', type=int);
    parser.add_argument('-o', '--out',   default='./output.pdf', dest='dout', help='Output document.', type=str, required=True);
    parser.add_argument('-s', '--simg',  default='./image.png', dest='simg', help='Signature image.', type=str, required=True);
    parser.add_argument('-w', '--width', default=150, dest='width', help='Width of signature image.', type=int);
    parser.add_argument('-x', '--margin-x', default=32, dest='x', help='Margin left of signature image.', type=int);
    parser.add_argument('-y', '--margin-y', default=32, dest='y', help='Margin bottom of signature image.', type=int);
    parser.add_argument('-e', '--height', default=150, dest='height', help='Height of signature image.', type=int);

    # BOTTOM_CENTER = 0x010;
    # BOTTOM_RIGHT  = 0x011;
    # BOTTOM_LEFT   = 0x012;
    # TOP_CENTER    = 0x020;
    # TOP_RIGHT     = 0x021;
    # TOP_LEFT      = 0x022;

    # CENTER_RIGHT  = 0x001;
    # CENTER_LEFT   = 0x002;
    # CENTER        = 0x000;

    parser.add_argument('--bottom-center', dest="bottom_center", action='store_true');
    parser.add_argument('--bottom-right', dest="bottom_right", action='store_true');
    parser.add_argument('--bottom-left', dest="bottom_left", action='store_true');
    parser.add_argument('--top-center', dest="top_center", action='store_true');
    parser.add_argument('--top-right', dest="top_right", action='store_true');
    parser.add_argument('--top-left', dest="top_left", action='store_true');

    parser.add_argument('--center-right', dest="center_right", action='store_true');
    parser.add_argument('--center-left', dest="center_left", action='store_true');
    parser.add_argument('--center', dest="center", action='store_true');

    args = parser.parse_args();
    dct  = dsg.DocType.PDF if args.doctype == 'pdf' else dsg.DocType.WORD;
    dim  = (args.width, args.height);
    pos  = (args.x, args.y);
    pgn  = args.page;

    rlpos = dsg.RelativePos.BOTTOM_RIGHT;
    if args.bottom_center:  rlpos = dsg.RelativePos.BOTTOM_CENTER;
    if args.bottom_right:   rlpos = dsg.RelativePos.BOTTOM_RIGHT;
    if args.bottom_left:    rlpos = dsg.RelativePos.BOTTOM_LEFT;
    if args.top_center:     rlpos = dsg.RelativePos.TOP_CENTER;
    if args.top_right:      rlpos = dsg.RelativePos.TOP_RIGHT;
    if args.top_left:       rlpos = dsg.RelativePos.TOP_LEFT;

    if args.center_right:   rlpos = dsg.RelativePos.CENTER_RIGHT;
    if args.center_left:    rlpos = dsg.RelativePos.CENTER_LEFT;
    if args.center:         rlpos = dsg.RelativePos.CENTER;

    # make signature
    dsg.sign(din=args.din, dout=args.dout, simg=args.simg, doctype=dct, pgn=pgn, sdim=dim, pos=pos, rlpos=rlpos);

